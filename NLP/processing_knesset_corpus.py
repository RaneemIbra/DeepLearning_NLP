import os
import re
import json
import sys
from docx import Document

# a function that divides sentences in the given text
def sentence_division(text):
    # define a list for all the sentences
    sentences = []
    # start from an empty sentence that will accumulate the chars
    sentence = ""
    # iterate over the whole text in order to see where to split
    for char in text:
        # start accumulating chars
        sentence += char
        # check if the char is equals to one of the sentence ending punctuation
        if char in ".!?:":
            # check if the sentence is longer than 1 to make sure it is a valid sentence and strip it from whitespaces
            if len(sentence.strip()) > 1:
                # append the sentence to the list
                sentences.append(sentence.strip())
                sentence = ""
    # one last check for the last white spaces
    if sentence.strip():
        sentences.append(sentence.strip())
    return sentences

# a funtion that validates the correctness of the sentence as a valid hebrew sentence
def validate_sentence(sentence):
    # we used regex to detect hebrew alphabet
    sentence_contains_hebrerw_alphabet = any("\u0590" <= char <= "\u05FF" for char in sentence)
    # if the sentence does not include hebrew alphabet then we want to exclude it
    if not sentence_contains_hebrerw_alphabet:
        return False
    # split the sentence to words separated by white spaces and remove any white spaces in the start or the end
    words = re.split(r'\s+', sentence.strip())
    # we use regex to check if the string contains only chars that are not alpha numeric values and certainly not hebrew letters
    if all(re.fullmatch(r"[^\w\u0590-\u05FF]+", word) for word in words):
        return False
    # validate that the sentence is complete
    if "..." in sentence or "---" in sentence:
        return False
    return True

# Function to tokenize a sentence into words and symbols
def Tokenize(sentence):
    # regex line to match hebrew letters
    hebrew_letters = r"[\u0590-\u05FF]+"
    # regex line to match hebrew punctuation
    hebrew_punctuation = r"[^\w\u0590-\u05FF\s]"
    # regex line to match number patterns
    number_pattern = r"[$₪€£]?\d+(?:\.\d+)?%?"
    # regex line to match url patterns
    url_pattern = r"https?://\S+|www\.\S+"
    # regex line to match email patterns
    email_pattern = r"[\w.+-]+@\w+\.\w+"
    # regex line to match parentheses
    parentheses_pattern = r"\([^()]*\)"
    # regex line to match hyphenated texts
    hyphenated_pattern = r"[\u0590-\u05FF]+(?:-[\u0590-\u05FF]+)*"
    # regex line to match abbriviations / initials (rashi tevot)
    abbreviation_pattern = r"\b(?:[\u0590-\u05FF]\.){2,}"
    # combine all the patterns to start the tokenization
    combined_pattern = f"({hebrew_letters}|{hebrew_punctuation}|{number_pattern}|{url_pattern}|{email_pattern}|{parentheses_pattern}|{hyphenated_pattern}|{abbreviation_pattern})"
    # now we tokenize based on all the patterns that we defined as well as whitespaces
    # that happen from behind the scenes in the findall function
    tokens = re.findall(combined_pattern, sentence)
    return [token for token in tokens if token.strip()]

# a function to extract the person who spoke the important text
def extract_speaker_name(text):
    # titles that might preceed the name of the speaker
    valid_title_keywords = [
        "יושב ראש", "חבר הכנסת", "חברת הכנסת", 
        "שר", "שרה", "ראש הממשלה", "נשיא", 
        "סגן", "מזכיר", "מזכירה", "יו\"ר"
    ]
    # regex to define the pattern we are looking for in order to extract the name of the speaker
    speaker_regex_pattern = r"^([\u0590-\u05FF\s\(\)\"'`]+):"
    # try to find the speaker in the text by matching it to the regex
    speaker_found = re.match(speaker_regex_pattern, text)
    if speaker_found:
        # take the uncleaned speaker name from the matched pattern, then strip it from the white spaces
        name_uncleaned = speaker_found.group(1).strip()
        # clean the name from parentheses :)
        clean_name = re.sub(r"\s*\(.*?\)", "", name_uncleaned)
        # clean the name from quotation marks
        clean_name = re.sub(r"\"|'", "", clean_name)
        # created a boolean variable to check if the name is as intended
        valid_name = (
            any(keyword in clean_name for keyword in valid_title_keywords) or
            bool(re.fullmatch(r"[\u0590-\u05FF\s]+", clean_name))
        )
        # If valid, return the cleaned name; otherwise, return "Unknown Speaker"
        return clean_name if valid_name else "Unknown Speaker"
    else:
        return "Unknown Speaker"
    
# a function to extract the number written in text (מספר ישיבה)
def hebrew_text_to_number(hebrew_number):
    # a dictionary mapping Hebrew words to integer equivalents
    hebrew_to_int = {
        "אחד": 1, "אחת": 1, "שניים": 2, "שתיים": 2, "שלושה": 3, "שלוש": 3, "ארבעה": 4, "ארבע": 4,
        "חמשה": 5, "חמש": 5, "ששה": 6, "שש": 6, "שבעה": 7, "שבע": 7, "שמונה": 8, "תשעה": 9, "תשע": 9,
        "עשרה": 10, "עשר": 10, "עשרים": 20, "שלושים": 30, "ארבעים": 40, "חמישים": 50,
        "שישים": 60, "שבעים": 70, "שמונים": 80, "תשעים": 90, "מאה": 100, "מאתיים": 200,
        "שלוש-מאות": 300, "ארבע-מאות": 400
    }
    # strip number from white spaces
    hebrew_number = hebrew_number.strip()
    # handle the prefix "ה"
    if hebrew_number.startswith("ה"):
        hebrew_number = hebrew_number[1:]
    # replace "-ו" and "ו" used as connectors with spaces
    hebrew_number = hebrew_number.replace("-ו", " ").replace("ו", " ")
    # split into components
    components = re.split(r"[-\s]", hebrew_number)
    # calculate the total value
    total = 0
    for word in components:
        value = hebrew_to_int.get(word, 0)
        if value == 0 and word:
            # for debugging issues
            print(f"Unrecognized Hebrew numeral: {word}")
        total += value
    return total



# function to start processing the files following the requirements of the HW
def file_processing_function(input_folder, output_file):
    global global_token_count
    # a line to read the docx files from the directory and store them in the list of files
    files = [file_name for file_name in os.listdir(input_folder) if file_name.endswith(".docx")]
    # print the files for debugging
    print(f"files: {files}")
    json_list = []
    # go over all the files to start processing them
    for file in files:
        # we have to match the file to the regex of the same format then we start extracting
        match = re.search(r'(\d+)_pt', file)
        # extract the kenest number from the matched file as it will the first number in the match
        # also added a fallback value just in case the code is being tested on other files
        keneset_number = int(match.group(1)) if match else -1
        # if the file name contains ptm, then it is a plenary protocol, else it is a committee
        protocol_type = "plenary" if "ptm" in file else "committee" if "ptv" in file else "undefined"
        protocol_number = None
        # a try except block just as requested to handle exceptions that could be caused
        # from opening files or other issues that may rise
        try:
            # since we have the path for the input and the file path we need to join them
            document_path = os.path.join(input_folder, file)
            # now we open the document to start working on it
            doc = Document(document_path)
            # since the protocol number is in the beginning of the documents
            for paragraph in doc.paragraphs:
                # try to match the protocol number with the regex
                protocol_number_match_regex = re.search(r"פרוטוקול מס'? (\d+)", paragraph.text)
                # if the number matched then we give the protocil_number variable the value
                # of the protocol number that appears in the document otherwise we assign
                # a fallback value which is -1 just as requested
                if protocol_number_match_regex:
                    protocol_number = int(protocol_number_match_regex.group(1))
                    break
                if "הישיבה" in paragraph.text:
                    hebrew_number_match = re.search(r"הישיבה\s([\u0590-\u05FF\-]+)", paragraph.text)
                    if hebrew_number_match:
                        protocol_number = hebrew_text_to_number(hebrew_number_match.group(1))
                        break
            if protocol_number is None or protocol_number == 0:
                protocol_number = -1
            # define a last speaker variable just in case we want to attribute parts of the speech to a speaker
            last_speaker = None
            # loop over the paragraph to start extracting the names of the speakers
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                # handle the case where we don't have a text
                if not text:
                    continue
                # call the function to extract the name, the problem from previously was with the language backticks
                speaker_name = extract_speaker_name(text)
                if speaker_name != "Unknown Speaker":
                    # if a valid speaker name is found, update the last speaker context
                    last_speaker = speaker_name
                    # we use slicing to get the text after the speaker name appears
                    spoken_text = text[len(speaker_name) + 1:].strip()
                    # divide the sentences following the logic implemented in the function
                    sentences = sentence_division(spoken_text)
                    # iterate over all the sentences to start processing them
                    for sentence in sentences:
                        if validate_sentence(sentence):
                            # call the function to tokenize the sentences
                            tokens = Tokenize(sentence)
                            # check if the tokens are more than 4 so that we can count them as sentences that
                            # we can work with them, and then we append the data to the list
                            if len(tokens) >= 4:
                                global_token_count += len(tokens)
                                json_list.append({
                                    "protocol_name": file,
                                    "knesset_number": keneset_number,
                                    "protocol_type": protocol_type,
                                    "protocol_number": protocol_number,
                                    "speaker_name": speaker_name,
                                    "sentence_text": " ".join(tokens)
                                })
                # if there was no speaker, we will want to attribute the text to the last speaker
                elif last_speaker:
                    # repeat the same process as before
                    additional_sentences = sentence_division(text)
                    for sentence in additional_sentences:
                        if validate_sentence(sentence):
                            tokens = Tokenize(sentence)
                            if len(tokens) >= 4:
                                global_token_count += len(tokens)
                                json_list.append({
                                    "protocol_name": file,
                                    "knesset_number": keneset_number,
                                    "protocol_type": protocol_type,
                                    "protocol_number": protocol_number,
                                    "speaker_name": last_speaker,
                                    "sentence_text": " ".join(tokens)
                                })
        except Exception as e:
            print(f"{file}: {e}")
    # here we open the json file with a permission to write and with the encoding that was requested
    # then we write each entry from the list that we saved to the file
    with open(output_file, "w", encoding="utf-8") as jsonl_file:
        for entry in json_list:
            jsonl_file.write(json.dumps(entry, ensure_ascii=False) + "\n")
    print(f"file saved at: {output_file}")
    print(f"tokens count: {global_token_count}")

# main entry for the program
if __name__ == "__main__":
    global_token_count = 0
    # check if the input taken from the user is in the correct format
    if len(sys.argv) != 3:
        sys.exit(1)
    # take the first argument as the folder path, and the second argument as the path
    # for where to save the output file
    input_folder = sys.argv[1]
    output_file = sys.argv[2]
    file_processing_function(input_folder, output_file)